"""
Format pandoc-generated docx:
- Font: Times New Roman (ASCII) + SimSun (East Asian), Size 小四 (12pt)
- Line spacing: 1.25
- Tables: three-line table style
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

INPUT_PATH = "outline_raw.docx"
OUTPUT_PATH = "outline.docx"

def set_run_font(run, ascii_font="Times New Roman", eastasia_font="宋体", size_pt=12):
    """Set ASCII and EastAsian fonts for a run."""
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    # East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), eastasia_font)
    # Also set hint for complex script if needed
    rFonts.set(qn('w:cs'), ascii_font)

def set_paragraph_format(paragraph, line_spacing=1.25, space_after=Pt(6)):
    """Set line spacing and space after."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = space_after

def set_three_line_table(table):
    """Convert table to three-line table style."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr {}></w:tblPr>'.format(nsdecls('w')))

    # Remove existing borders
    existing_borders = tblPr.xpath('./w:tblBorders')
    for b in existing_borders:
        tblPr.remove(b)

    # Add three-line borders: top, header-bottom, bottom
    # Use 1.5pt for top/bottom, 0.5pt for header bottom (middle horizontal)
    borders_xml = (
        '<w:tblBorders {}>'
        '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    ).format(nsdecls('w'))
    tblPr.append(parse_xml(borders_xml))

    # Add header row bottom border (insideH for first row) by setting bottom border of first row cells
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Remove existing cell borders
            existing = tcPr.xpath('./w:tcBorders')
            for e in existing:
                tcPr.remove(e)
            tcBorders = parse_xml(
                '<w:tcBorders {}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'.format(nsdecls('w'))
            )
            tcPr.append(tcBorders)

    # Remove vertical borders for all other cells
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            continue
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            existing = tcPr.xpath('./w:tcBorders')
            for e in existing:
                tcPr.remove(e)
            tcBorders = parse_xml(
                '<w:tcBorders {}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>'.format(nsdecls('w'))
            )
            tcPr.append(tcBorders)

def main():
    doc = Document(INPUT_PATH)

    # Set default document font via Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.element.rPr.rFonts.set(qn('w:cs'), 'Times New Roman')

    # Set default paragraph format
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)

    # Also adjust heading styles
    for level in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.size = Pt(14 - level)  # H1=14, H2=13, H3=12...
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            heading_style.paragraph_format.line_spacing = 1.25
            heading_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    # Process all paragraphs
    for paragraph in doc.paragraphs:
        set_paragraph_format(paragraph)
        # Determine font size based on heading level
        style_name = paragraph.style.name
        if style_name == 'Heading 1':
            size_pt = 16
            bold = True
        elif style_name == 'Heading 2':
            size_pt = 14
            bold = True
        elif style_name == 'Heading 3':
            size_pt = 12
            bold = True
        elif style_name == 'Heading 4':
            size_pt = 12
            bold = True
        else:
            size_pt = 12
            bold = False

        for run in paragraph.runs:
            set_run_font(run, size_pt=size_pt)
            if bold:
                run.font.bold = True

    # Process all tables
    for table in doc.tables:
        set_three_line_table(table)
        # Also set font for text inside tables
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph)
                    for run in paragraph.runs:
                        set_run_font(run)

    doc.save(OUTPUT_PATH)
    print(f"Saved formatted document to {OUTPUT_PATH}")

if __name__ == '__main__':
    main()
